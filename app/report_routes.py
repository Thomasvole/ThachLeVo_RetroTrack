# app/report_routes.py

import io
import os

import xlsxwriter
from docx import Document
from docx.shared import Inches
from flask import (
    render_template, redirect, url_for,
    flash, session, send_file
)
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Table, TableStyle, Image
)

from . import app
from .helpers import generate_summary
from .models import User

# Register custom fonts for PDF output (Arial and Arial-Bold)
pdfmetrics.registerFont(TTFont('Arial', r'C:\Windows\Fonts\arial.ttf'))
pdfmetrics.registerFont(TTFont('Arial-Bold', r'C:\Windows\Fonts\arialbd.ttf'))


@app.route('/report/<int:file_id>')
def report_preview(file_id):
    """Render an HTML preview of the report before download."""
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)
    chart_url = f"/static/plots/cost_waste_{file_id}.png"
    return render_template(
        "report_preview.html",
        summary=summary,
        file_id=file_id,
        chart_url=chart_url
    )


@app.route('/download_report_pdf/<int:file_id>')
def download_report_pdf(file_id):
    """Build and send the report as a PDF file."""
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    # Create an in-memory buffer and SimpleDocTemplate for PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=20, rightMargin=20,
        topMargin=20, bottomMargin=20
    )

    # Override default styles to use registered Arial fonts
    styles = getSampleStyleSheet()
    styles['Title'].fontName = 'Arial-Bold'
    styles['Heading2'].fontName = 'Arial-Bold'
    styles['Heading5'].fontName = 'Arial-Bold'
    styles['BodyText'].fontName = 'Arial'

    # Define paragraph styles for body and headers
    body_style = ParagraphStyle(
        'body', parent=styles['BodyText'],
        fontName='Arial', fontSize=8, leading=10
    )
    header_style = ParagraphStyle(
        'hdr', parent=styles['Heading5'],
        fontName='Arial-Bold', fontSize=8,
        leading=10, alignment=1  # centered
    )

    elems = []
    # Title and spacing
    elems.append(Paragraph("RetroTrack Logistics Inefficiency Report", styles['Title']))
    elems.append(Spacer(1, 12))

    # --- Inefficient Routes ---
    if summary['ineff_table']:
        keys = [
            "base_address", "shipping_address",
            "starting_time", "expected_delivery_time",
            "actual_delivery_time", "delay_hours",
            "optimized_delivery_time", "time_saved"
        ]
        headers = [
            "Base Address", "Shipping Address",
            "Start Time", "Expected Time",
            "Actual Time", "Delay (h)",
            "Optimized (h)", "Time Saved (h)"
        ]

        # Build table data: header row + one row per route
        data = [[Paragraph(h, header_style) for h in headers]]
        for row in summary['ineff_table']:
            data.append([Paragraph(str(row[k]), body_style) for k in keys])

        # Create and style the table
        col_widths = [120, 120, 60, 60, 60, 40, 40, 40]
        tbl = Table(data, repeatRows=1, colWidths=col_widths)
        tbl.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elems.append(Paragraph("Inefficient Routes", styles['Heading2']))
        elems.append(tbl)
        elems.append(Spacer(1, 24))

    # --- Cost Analysis ---
    if summary['cost_table']:
        keys2 = [
            "route_id", "base_address", "shipping_address",
            "actual_duration", "optimized_time",
            "max_delivery_cost", "optimized_cost",
            "actual_cost", "cost_saved"
        ]
        headers2 = [
            "Route ID", "Base Address", "Shipping Address",
            "Actual Dur (h)", "Optimized (h)",
            "Max Cost", "Opt Cost", "Act Cost", "Saved"
        ]

        data2 = [[Paragraph(h, header_style) for h in headers2]]
        for row in summary['cost_table']:
            data2.append([Paragraph(str(row[k]), body_style) for k in keys2])

        col_widths2 = [30, 90, 90, 50, 50, 60, 60, 60, 60]
        tbl2 = Table(data2, repeatRows=1, colWidths=col_widths2)
        tbl2.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elems.append(Paragraph("Cost Analysis", styles['Heading2']))
        elems.append(tbl2)
        elems.append(Spacer(1, 24))

    # --- Chart ---
    chart_file = os.path.join(
        app.root_path, 'static', 'plots',
        f"cost_waste_{file_id}.png"
    )
    if os.path.exists(chart_file):
        elems.append(Paragraph("Cost Saved Visualization", styles['Heading2']))
        elems.append(Spacer(1, 12))
        elems.append(Image(chart_file, width=400, height=200))
        elems.append(Spacer(1, 24))

    # Build PDF and send as attachment
    doc.build(elems)
    buffer.seek(0)
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'report_{file_id}.pdf'
    )


@app.route('/download_report_word/<int:file_id>')
def download_report_word(file_id):
    """Build and send the report as a Word document."""

    # Ensure the user is logged in
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    # Fetch current user and generate the summary data
    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    # Create a new Word document
    docx = Document()
    docx.add_heading("RetroTrack Logistics Inefficiency Report", 0)

    # --- Summary Table (2 columns) ---
    tbl = docx.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for key in [
        'file_name', 'upload_date', 'inefficient_routes',
        'total_delayed_hours', 'avg_delayed_hours',
        'total_time_saved', 'avg_time_saved',
        'total_cost_saved', 'avg_cost_saved'
    ]:
        row = tbl.add_row().cells
        # Capitalize and format metric name
        row[0].text = key.replace('_', ' ').title()
        # Convert the summary value to string
        row[1].text = str(summary[key])
    docx.add_paragraph()  # Add spacing after the table

    # --- Inefficient Routes Table ---
    if summary['ineff_table']:
        keys = [
            "base_address", "shipping_address",
            "starting_time", "expected_delivery_time",
            "actual_delivery_time", "delay_hours",
            "optimized_delivery_time", "time_saved"
        ]
        hdrs = [
            "Base Address", "Shipping Address",
            "Start Time", "Expected Time",
            "Actual Time", "Delay (h)",
            "Optimized (h)", "Time Saved (h)"
        ]

        # Create table with header row
        t = docx.add_table(rows=1, cols=len(keys))
        for i, h in enumerate(hdrs):
            t.rows[0].cells[i].text = h
        # Populate rows with data
        for row in summary['ineff_table']:
            r = t.add_row().cells
            for i, k in enumerate(keys):
                r[i].text = str(row[k])
        docx.add_paragraph()

    # --- Cost Analysis Table ---
    if summary['cost_table']:
        keys = [
            "route_id", "base_address", "shipping_address",
            "actual_duration", "optimized_time",
            "max_delivery_cost", "optimized_cost",
            "actual_cost", "cost_saved"
        ]
        hdrs = [
            "Route ID", "Base Address", "Shipping Address",
            "Actual Dur (h)", "Optimized (h)",
            "Max Cost", "Opt Cost", "Act Cost", "Saved"
        ]
        t2 = docx.add_table(rows=1, cols=len(keys))
        for i, h in enumerate(hdrs):
            t2.rows[0].cells[i].text = h
        for row in summary['cost_table']:
            r = t2.add_row().cells
            for i, k in enumerate(keys):
                r[i].text = str(row[k])

    # --- Insert Chart if Available ---
    chart_path = os.path.join(
        app.root_path, 'static', 'plots',
        f"cost_waste_{file_id}.png"
    )
    if os.path.exists(chart_path):
        # Embed the chart image at 6 inches wide
        docx.add_picture(chart_path, width=Inches(6))

    # Save document to buffer and send as attachment
    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"report_{file_id}.docx"
    )


@app.route('/download_report_excel/<int:file_id>')
def download_report_excel(file_id):
    """
       Build and send the report as an Excel (.xlsx) workbook.
    """
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    # Fetch current user and generate the summary data
    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    # Create an in-memory Excel workbook
    buf = io.BytesIO()
    wb = xlsxwriter.Workbook(buf)

    # Add worksheets for summary, inefficient routes, and cost analysis
    summary_ws = wb.add_worksheet("Summary")
    ineff_ws = wb.add_worksheet("Inefficient Routes")
    cost_ws = wb.add_worksheet("Cost Analysis")

    # --- Write Summary ---
    row = 0
    for key in [
        'file_name', 'upload_date', 'user_name', 'user_email',
        'inefficient_routes', 'total_delayed_hours',
        'avg_delayed_hours', 'total_time_saved',
        'avg_time_saved', 'total_cost_saved', 'avg_cost_saved'
    ]:
        # Metric name in first column
        summary_ws.write(row, 0, key.replace('_', ' ').title())
        # Metric value in second column
        summary_ws.write(row, 1, summary[key])
        row += 1

    # --- Write Inefficient Routes Table ---
    headers = list(summary['ineff_table'][0].keys()) if summary['ineff_table'] else []
    for col, h in enumerate(headers):
        ineff_ws.write(0, col, h)
    for i, item in enumerate(summary['ineff_table'], 1):
        for j, h in enumerate(headers):
            ineff_ws.write(i, j, item[h])

    # --- Write Cost Analysis Table ---
    headers = list(summary['cost_table'][0].keys()) if summary['cost_table'] else []
    for col, h in enumerate(headers):
        cost_ws.write(0, col, h)
    for i, item in enumerate(summary['cost_table'], 1):
        for j, h in enumerate(headers):
            cost_ws.write(i, j, item[h])

    # --- Insert Chart Image if Present ---
    chart_path = os.path.join(
        app.root_path, 'static', 'plots',
        f"cost_waste_{file_id}.png"
    )
    if os.path.exists(chart_path):
        # Place the chart at cell N2
        cost_ws.insert_image('N2', chart_path)

    # Finalize and close the workbook, then return it as a download
    wb.close()
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f"report_{file_id}.xlsx"
    )
