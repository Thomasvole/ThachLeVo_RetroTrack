import io
import json
import os
from datetime import datetime, timedelta
from urllib.parse import quote_plus

import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import pandas as pd
import requests  # For external API calls
import xlsxwriter
from docx import Document
from docx.shared import Inches
from flask import Flask, render_template, redirect, url_for, flash, session, request
from flask import send_file
from flask_bcrypt import Bcrypt
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import FlaskForm
from fpdf import FPDF
from werkzeug.utils import secure_filename
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, EqualTo

#########################################################################
# Initialize Flask app and configuration
#########################################################################
app = Flask(__name__)
app.config['SECRET_KEY'] = 'a_very_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx', 'xls', 'xlsx'}
# Use your Geoapify API key (ensure the key is active and not restricted)
app.config['GEOAPIFY_API_KEY'] = "380c58ce52a64969a8f5a6e5ea6106da"

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)


#########################################################################
# Database Models
#########################################################################
class User(db.Model):
    """Represents a user in the system."""
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100), nullable=False)
    last_name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)

    def __repr__(self):
        return f"User('{self.email}')"


class File(db.Model):
    """Represents a file uploaded by a user."""
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    upload_date = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    size = db.Column(db.Float, nullable=False)
    user_id = db.Column(db.Integer, nullable=False)  # storing owner user id
    parsed_data = db.Column(db.Text, nullable=True)

    def __repr__(self):
        return f"File('{self.filename}', UserID={self.user_id})"


class InefficientRoute(db.Model):
    """
    Stores route information when the actual delivery time lags the expected delivery time by >24 hours.
    Only rows with complete data in these columns are stored:
      - Base Address, Shipping Address, Starting Time,
      - Expected Delivery Time, Actual Delivery Time,
      - Expected Delivery Cost, Actual Delivery Cost, Max Delivery Cost.
    (Delay is computed on the fly when displaying.)
    Additionally, we now store the optimized delivery time (in hours) and time saved (in hours)
    so that these do not need to be recalculated on every page load.
    """
    id = db.Column(db.Integer, primary_key=True)
    file_id = db.Column(db.Integer, db.ForeignKey('file.id'), nullable=False)
    base_address = db.Column(db.String(255), nullable=False)
    shipping_address = db.Column(db.String(255), nullable=False)
    starting_time = db.Column(db.DateTime, nullable=False)
    expected_delivery_time = db.Column(db.DateTime, nullable=False)
    actual_delivery_time = db.Column(db.DateTime, nullable=False)
    expected_delivery_cost = db.Column(db.Float, nullable=False)
    actual_delivery_cost = db.Column(db.Float, nullable=False)
    max_delivery_cost = db.Column(db.Float, nullable=False)
    # New columns for caching computed optimized route info:
    optimized_delivery_time = db.Column(db.Float, nullable=True)
    time_saved = db.Column(db.Float, nullable=True)

    def __repr__(self):
        delay = (self.actual_delivery_time - self.expected_delivery_time).total_seconds() / 3600.0
        return f"InefficientRoute(FileID={self.file_id}, BaseAddress={self.base_address}, Delay={round(delay, 2)}h)"


#########################################################################
# Forms
#########################################################################
class RegistrationForm(FlaskForm):
    first_name = StringField('First Name', validators=[DataRequired()])
    last_name = StringField('Last Name', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Register')


class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Login')


class ProfileEditForm(FlaskForm):
    first_name = StringField('First Name', validators=[DataRequired()])
    last_name = StringField('Last Name', validators=[DataRequired()])
    submit = SubmitField('Save Changes')


#########################################################################
# Helper Functions
#########################################################################
def allowed_file(filename):
    """Return True if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def to_float(value):
    """Convert a value (possibly containing commas) to a float."""
    try:
        return float(str(value).replace(',', '').strip())
    except Exception:
        return None


def parse_excel(file_path):
    """
    Extracts inefficient route data from an Excel file.
    Expected columns:
      - Base Address, Shipping Address, Starting Time,
      - Expected Delivery Time (hours), Actual Delivery Time (hours),
      - Expected Delivery Cost (VND), Actual Delivery Cost (VND),
      - Max Delivery Cost (VND/hr)
    Converts numeric hours to datetime objects by adding them to the Starting Time.
    Keeps rows where delay (act_hours - exp_hours) > 24.
    Returns a dictionary with key "inefficient_routes".
    """
    required_cols = [
        "Base Address",
        "Shipping Address",
        "Starting Time",
        "Expected Delivery Time (hours)",
        "Actual Delivery Time (hours)",
        "Expected Delivery Cost (VND)",
        "Actual Delivery Cost (VND)",
        "Max Delivery Cost (VND/hr)"
    ]
    result = {"inefficient_routes": []}
    try:
        xls = pd.ExcelFile(file_path)
        for sheet in xls.sheet_names:
            df = xls.parse(sheet)
            df.columns = df.columns.str.strip()
            if not set(required_cols).issubset(df.columns):
                continue
            for _, row in df.iterrows():
                if any(pd.isnull(row[col]) for col in required_cols):
                    continue
                st = row["Starting Time"]
                exp_hours = to_float(row["Expected Delivery Time (hours)"])
                act_hours = to_float(row["Actual Delivery Time (hours)"])
                if exp_hours is None or act_hours is None:
                    continue
                diff = act_hours - exp_hours
                if diff > 24:
                    expected_dt = st + timedelta(hours=exp_hours)
                    actual_dt = st + timedelta(hours=act_hours)
                    ec = to_float(row["Expected Delivery Cost (VND)"])
                    ac = to_float(row["Actual Delivery Cost (VND)"])
                    mc = to_float(row["Max Delivery Cost (VND/hr)"])
                    if None in (ec, ac, mc):
                        continue
                    route = {
                        "base_address": str(row["Base Address"]),
                        "shipping_address": str(row["Shipping Address"]),
                        "starting_time": st,
                        "expected_delivery_time": expected_dt,
                        "actual_delivery_time": actual_dt,
                        "expected_delivery_cost": ec,
                        "actual_delivery_cost": ac,
                        "max_delivery_cost": mc,
                        "delay_hours": diff
                    }
                    result["inefficient_routes"].append(route)
        print("Parsed Inefficient Routes:")
        print(json.dumps(result, default=str, indent=4, ensure_ascii=False))
        return result
    except Exception as e:
        print(f"Error processing Excel file: {e}")
        return result


# ------------------------------
# Geoapify Helper Functions
# ------------------------------
def get_coordinates(address):
    """
    Returns the (latitude, longitude) of the given address using Geoapify's Forward Geocoding API.
    Uses URL encoding, sets language to Vietnamese, and requests JSON output.
    """
    try:
        encoded_address = quote_plus(address)
        geocode_url = (
            f"https://api.geoapify.com/v1/geocode/search?text={encoded_address}"
            f"&apiKey={app.config['GEOAPIFY_API_KEY']}&lang=vi&limit=1&format=json"
        )
        print("Geocode URL:", geocode_url)
        headers = {"Accept": "application/json", "User-Agent": "FlaskGeoapifyClient/1.0"}
        response = requests.get(geocode_url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            print("Geocode response data:", data)
            if data.get("results") and len(data["results"]) > 0:
                result = data["results"][0]
                return result.get("lat"), result.get("lon")
            else:
                print("No geocoding results found for address:", address)
        else:
            print("Geocode request failed with status:", response.status_code)
    except Exception as e:
        print("Error geocoding address", address, e)
    return None


def get_optimized_route_time(base_address, shipping_address):
    """
    Returns the optimized route travel time (in hours) between base_address and shipping_address using Geoapify's Routing API.
    Uses the coordinates retrieved via the Forward Geocoding API.
    """
    start_coords = get_coordinates(base_address)
    end_coords = get_coordinates(shipping_address)
    if not start_coords or not end_coords:
        print("Failed to get coordinates for either address.")
        return None

    routing_url = (
        f"https://api.geoapify.com/v1/routing?"
        f"waypoints={start_coords[0]},{start_coords[1]}|{end_coords[0]},{end_coords[1]}"
        f"&mode=drive&type=short&units=metric&apiKey={app.config['GEOAPIFY_API_KEY']}&limit=1&format=json"
    )
    print("Routing URL:", routing_url)
    headers = {"Accept": "application/json", "User-Agent": "FlaskGeoapifyClient/1.0"}
    try:
        response = requests.get(routing_url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            print("Routing response data:", data)
            if data.get("results") and len(data["results"]) > 0:
                result = data["results"][0]
                travel_time_seconds = result.get("time")
                if travel_time_seconds is not None:
                    return travel_time_seconds / 3600.0
            else:
                print("No routing results returned.")
        else:
            print("Routing request failed with status:", response.status_code)
    except Exception as e:
        print("Error fetching optimized route from Geoapify", e)
    return None


# --- New helper function to update cached optimized info ---
def update_cached_routes(routes):
    """
    For each route in the given list that does not have cached optimized_delivery_time or time_saved,
    compute these values using the Routing API and update the database.
    """
    updated = False
    for r in routes:
        if r.optimized_delivery_time is None or r.time_saved is None:
            optimized_time = get_optimized_route_time(r.base_address, r.shipping_address)
            if optimized_time is not None:
                actual_duration = (r.actual_delivery_time - r.starting_time).total_seconds() / 3600.0
                if actual_duration > optimized_time:
                    r.optimized_delivery_time = round(optimized_time, 2)
                    r.time_saved = round(actual_duration - optimized_time, 2)
                    updated = True
    if updated:
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print("Error updating cached optimized info:", e)


#########################################################################
# Context Processor
#########################################################################
@app.context_processor
def inject_now():
    return {"now": datetime.utcnow()}


#########################################################################
# Routes
#########################################################################
@app.route('/')
def home():
    user = None
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        if user:
            user_files = File.query.filter_by(user_id=user.id).all()
            session['user_files'] = [{"id": f.id, "filename": f.filename} for f in user_files]
    return render_template('homepage.html', user=user)


@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    if form.validate_on_submit():
        hashed = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        usr = User(
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            email=form.email.data,
            password=hashed
        )
        db.session.add(usr)
        db.session.commit()
        flash("Account created. Please log in.", "success")
        return redirect(url_for('login'))
    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        usr = User.query.filter_by(email=form.email.data).first()
        if usr and bcrypt.check_password_hash(usr.password, form.password.data):
            session['user_id'] = usr.id
            session['user_name'] = f"{usr.first_name} {usr.last_name}"
            session['user_email'] = usr.email
            flash("Logged in successfully.", "success")
            return redirect(url_for('home'))
        else:
            flash("Login failed. Check email/password.", "danger")
    return render_template('login.html', form=form)


@app.route('/logout')
def logout():
    session.clear()
    flash("Logged out.", "success")
    return redirect(url_for('home'))


@app.route('/profile')
def profile():
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))
    usr = User.query.get(session['user_id'])
    if not usr:
        flash("User not found.", "danger")
        return redirect(url_for('logout'))
    return render_template('profile.html', user=usr)


@app.route('/edit-profile', methods=['GET', 'POST'])
def edit_profile():
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))
    usr = User.query.get(session['user_id'])
    if not usr:
        flash("User not found.", "danger")
        return redirect(url_for('logout'))
    form = ProfileEditForm(obj=usr)
    if form.validate_on_submit():
        usr.first_name = form.first_name.data
        usr.last_name = form.last_name.data
        db.session.commit()
        flash("Profile updated.", "success")
        return redirect(url_for('profile'))
    return render_template('edit_profile.html', form=form, user=usr)


@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if 'user_id' not in session:
        flash("Login required to upload files.", "danger")
        return redirect(url_for('login'))

    if request.method == 'POST':
        if request.content_length and request.content_length > 10 * 1024 * 1024:
            flash("File too large (>10MB).", "danger")
            return redirect(request.url)

        file = request.files.get('file')
        if not file or file.filename == '':
            flash("No file selected.", "danger")
            return redirect(request.url)

        if allowed_file(file.filename):
            fname = secure_filename(file.filename)
            fpath = os.path.join(app.config['UPLOAD_FOLDER'], fname)
            file.save(fpath)
            size_kb = os.path.getsize(fpath) / 1024.0

            if size_kb > 10240:
                os.remove(fpath)
                flash("File too large (>10MB).", "danger")
                return redirect(request.url)

            try:
                new_file = File(filename=fname, size=size_kb, user_id=session['user_id'])
                db.session.add(new_file)
                db.session.commit()  # new_file.id is available

                ext = fname.rsplit('.', 1)[1].lower()
                if ext in ['xls', 'xlsx']:
                    parsed = parse_excel(fpath)
                else:
                    flash("Only Excel files are supported for parsing inefficient routes.", "danger")
                    return redirect(url_for('upload'))

                new_file.parsed_data = json.dumps(parsed, default=str)
                db.session.commit()

                # Insert inefficient route records (only if delay > 24 hours)
                for route in parsed.get('inefficient_routes', []):
                    delay = route.get("delay_hours")
                    if delay is None:
                        ev = to_float(route["expected_delivery_time"])
                        av = to_float(route["actual_delivery_time"])
                        if ev is not None and av is not None:
                            delay = av - ev
                        else:
                            delay = (route["actual_delivery_time"] - route[
                                "expected_delivery_time"]).total_seconds() / 3600.0

                    if delay > 24:
                        ir = InefficientRoute(
                            file_id=new_file.id,
                            base_address=route["base_address"],
                            shipping_address=route["shipping_address"],
                            starting_time=route["starting_time"],
                            expected_delivery_time=route["expected_delivery_time"],
                            actual_delivery_time=route["actual_delivery_time"],
                            expected_delivery_cost=to_float(route["expected_delivery_cost"]),
                            actual_delivery_cost=to_float(route["actual_delivery_cost"]),
                            max_delivery_cost=to_float(route["max_delivery_cost"])
                        )
                        db.session.add(ir)
                db.session.commit()

                # Debug: Print inefficient routes to console.
                ineffs = InefficientRoute.query.filter_by(file_id=new_file.id).all()
                print("\n=== Inefficient Routes Stored in DB ===")
                print("Number of inefficient routes:", len(ineffs))
                for ir in ineffs:
                    computed_delay = (ir.actual_delivery_time - ir.expected_delivery_time).total_seconds() / 3600.0
                    print(f"FileID: {ir.file_id}, Base: {ir.base_address}, Delay: {round(computed_delay, 2)}h")
                print("=== End ===\n")

                flash("File uploaded and processed successfully!", "success")
            except Exception as e:
                flash("Error processing the file.", "danger")
                print("Error:", e)
            finally:
                if os.path.exists(fpath):
                    os.remove(fpath)

            return redirect(url_for('files'))
        else:
            flash("Invalid file extension.", "danger")
            return redirect(request.url)

    return render_template('upload.html')


@app.route('/files')
def files():
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))
    user_files = File.query.filter_by(user_id=session['user_id']).all()
    return render_template('files.html', files=user_files)


@app.route('/delete-file/<int:file_id>', methods=['POST'])
def delete_file(file_id):
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    this_file = File.query.get_or_404(file_id)

    if this_file.user_id != session['user_id']:
        flash("Not authorized.", "danger")
        return redirect(url_for('files'))

    fpath = os.path.join(app.config['UPLOAD_FOLDER'], this_file.filename)
    try:
        if os.path.exists(fpath):
            os.remove(fpath)
            flash("File successfully deleted.", "success")
        else:
            flash("File not found on server.", "warning")
    except Exception as e:
        flash(f"Error deleting file: {str(e)}", "danger")
        return redirect(url_for('files'))

    try:
        InefficientRoute.query.filter_by(file_id=file_id).delete()
        db.session.delete(this_file)
        db.session.commit()
        flash("File and related data deleted from the database.", "success")
    except Exception as e:
        flash(f"Error deleting file record: {str(e)}", "danger")

    return redirect(url_for('files'))


# --- Updated Inefficient Routes (List & Detail) Route ---
def show_inefficient(file_id):
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    # If POST: update the inefficient routes (and recalc optimized info)
    if request.method == 'POST':
        new_routes_count = 0
        user_files = File.query.filter_by(user_id=session['user_id']).all()
        for file in user_files:
            if file.parsed_data:
                try:
                    data = json.loads(file.parsed_data)
                    for route in data.get("inefficient_routes", []):
                        existing = InefficientRoute.query.filter_by(
                            file_id=file.id,
                            base_address=route["base_address"],
                            starting_time=route["starting_time"]
                        ).first()
                        if existing:
                            continue

                        delay = route.get("delay_hours")
                        if delay is None:
                            ev = to_float(route["expected_delivery_time"])
                            av = to_float(route["actual_delivery_time"])
                            if ev is not None and av is not None:
                                delay = av - ev
                            else:
                                continue

                        if delay > 24:
                            if isinstance(route["starting_time"], str):
                                try:
                                    start_time = datetime.fromisoformat(route["starting_time"])
                                except Exception:
                                    start_time = datetime.strptime(route["starting_time"], "%Y-%m-%d %H:%M:%S")
                            else:
                                start_time = route["starting_time"]

                            if isinstance(route["expected_delivery_time"], float):
                                expected_time = start_time + timedelta(hours=route["expected_delivery_time"])
                            elif isinstance(route["expected_delivery_time"], str):
                                try:
                                    expected_time = datetime.fromisoformat(route["expected_delivery_time"])
                                except Exception:
                                    expected_time = datetime.strptime(route["expected_delivery_time"],
                                                                      "%Y-%m-%d %H:%M:%S")
                            else:
                                expected_time = route["expected_delivery_time"]

                            if isinstance(route["actual_delivery_time"], float):
                                actual_time = start_time + timedelta(hours=route["actual_delivery_time"])
                            elif isinstance(route["actual_delivery_time"], str):
                                try:
                                    actual_time = datetime.fromisoformat(route["actual_delivery_time"])
                                except Exception:
                                    actual_time = datetime.strptime(route["actual_delivery_time"], "%Y-%m-%d %H:%M:%S")
                            else:
                                actual_time = route["actual_delivery_time"]

                            new_ir = InefficientRoute(
                                file_id=file.id,
                                base_address=route["base_address"],
                                shipping_address=route["shipping_address"],
                                starting_time=start_time,
                                expected_delivery_time=expected_time,
                                actual_delivery_time=actual_time,
                                expected_delivery_cost=to_float(route["expected_delivery_cost"]),
                                actual_delivery_cost=to_float(route["actual_delivery_cost"]),
                                max_delivery_cost=to_float(route["max_delivery_cost"])
                            )
                            db.session.add(new_ir)
                            new_routes_count += 1
                except Exception as e:
                    print("Error processing file", file.filename, e)

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            flash("An error occurred while saving to the database.", "danger")
            print("Commit error:", e)
            return redirect(url_for('show_inefficient'))

        if new_routes_count > 0:
            flash(f"{new_routes_count} inefficient route(s) identified and saved into the database.", "success")
        else:
            flash("No new inefficient routes found.", "info")
        return redirect(url_for('show_inefficient'))

    # If file_id is not provided, show a list of files that have inefficient routes.
    if file_id is None:
        file_ids = [fid for (fid,) in db.session.query(InefficientRoute.file_id).distinct().all()]
        files_list = File.query.filter(File.id.in_(file_ids), File.user_id == session['user_id']).all()
        return render_template("inefficient.html", files=files_list, routes=None, selected_file=None)
    else:
        routes = InefficientRoute.query.filter_by(file_id=file_id).all()
        update_cached_routes(routes)

        ineff_data = []
        for r in routes:
            try:
                delay = (r.actual_delivery_time - r.expected_delivery_time).total_seconds() / 3600.0
            except Exception:
                delay = "N/A"
            ineff_data.append({
                "file_id": r.file_id,
                "base_address": r.base_address,
                "shipping_address": r.shipping_address,
                "starting_time": r.starting_time.strftime("%Y-%m-%d %H:%M") if isinstance(r.starting_time,
                                                                                          datetime) else r.starting_time,
                "expected_delivery_time": r.expected_delivery_time.strftime("%Y-%m-%d %H:%M") if isinstance(
                    r.expected_delivery_time, datetime) else r.expected_delivery_time,
                "actual_delivery_time": r.actual_delivery_time.strftime("%Y-%m-%d %H:%M") if isinstance(
                    r.actual_delivery_time, datetime) else r.actual_delivery_time,
                "expected_delivery_cost": r.expected_delivery_cost,
                "actual_delivery_cost": r.actual_delivery_cost,
                "max_delivery_cost": r.max_delivery_cost,
                "delay_hours": round(delay, 2) if isinstance(delay, (int, float)) else delay,
                "optimized_delivery_time": r.optimized_delivery_time if r.optimized_delivery_time is not None else "N/A",
                "time_saved": r.time_saved if r.time_saved is not None else "N/A"
            })

        selected_file = File.query.get(file_id)
        return render_template("inefficient.html", files=None, routes=ineff_data, selected_file=selected_file)


def generate_cost_waste_chart(cost_data, file_id):
    """
    Generates and saves a cleaner bar chart of cost_saved for each route.
    """
    plot_dir = os.path.join('static', 'plots')
    os.makedirs(plot_dir, exist_ok=True)

    labels = [str(item['route_id']) for item in cost_data]
    waste_millions = [item['cost_saved'] / 1_000_000 for item in cost_data]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(labels, waste_millions)
    ax.set_xlabel('Route ID')
    ax.set_ylabel('Cost Saved (Million VND)')
    ax.set_title('Cost Saved per Inefficient Route')

    # Rotate and space x-ticks nicely
    ax.set_xticks(range(0, len(labels), 5))  # show every 5th label
    ax.set_xticklabels([labels[i] for i in range(0, len(labels), 5)], rotation=45, ha='right')
    ax.yaxis.set_major_formatter(mticker.FormatStrFormatter('%.1f'))

    plt.tight_layout()
    plot_filename = f"cost_waste_{file_id}.png"
    plot_path = os.path.join(plot_dir, plot_filename)
    plt.savefig(plot_path)
    plt.close(fig)

    return f"/static/plots/{plot_filename}"


def cost_analysis(file_id):
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    # 1) No file selected yet: list all files
    if file_id is None:
        files_list = File.query.filter_by(user_id=session['user_id']).all()
        return render_template("cost_analysis.html",
                               files=files_list,
                               cost_data=None,
                               chart_url=None,
                               selected_file=None)

    # 2) Build cost_data for the selected file
    routes = InefficientRoute.query.filter_by(file_id=file_id).all()
    update_cached_routes(routes)
    cost_data = []
    for r in routes:
        try:
            actual_duration = (r.actual_delivery_time - r.starting_time).total_seconds() / 3600.0
        except Exception:
            continue
        if r.optimized_delivery_time is None:
            continue

        optimized_time = r.optimized_delivery_time
        optimized_cost = r.max_delivery_cost * optimized_time
        actual_cost = r.max_delivery_cost * actual_duration
        cost_saved = actual_cost - optimized_cost

        cost_data.append({
            "route_id": r.id,
            "base_address": r.base_address,
            "shipping_address": r.shipping_address,
            "actual_duration": round(actual_duration, 2),
            "optimized_time": round(optimized_time, 2),
            "max_delivery_cost": r.max_delivery_cost,
            "optimized_cost": round(optimized_cost, 2),
            "actual_cost": round(actual_cost, 2),
            "cost_saved": round(cost_saved, 2)
        })

    # 3) Generate chart if data exists
    chart_url = None
    if cost_data:
        chart_url = generate_cost_waste_chart(cost_data, file_id)

    # 4) Fetch the File record so template can show name + id
    selected_file = File.query.get_or_404(file_id)

    return render_template("cost_analysis.html",
                           files=None,
                           cost_data=cost_data,
                           chart_url=chart_url,
                           selected_file=selected_file)


def generate_summary(file_id, user):
    file = File.query.get_or_404(file_id)
    routes = InefficientRoute.query.filter_by(file_id=file_id).all()
    update_cached_routes(routes)

    inefficient_routes = len([
        r for r in routes
        if (r.actual_delivery_time - r.expected_delivery_time).total_seconds() / 3600 > 24
    ])

    total_delayed_hours = sum(
        (r.actual_delivery_time - r.expected_delivery_time).total_seconds() / 3600 for r in routes
    )
    avg_delayed_hours = total_delayed_hours / inefficient_routes if inefficient_routes else 0

    total_time_saved = sum(r.time_saved for r in routes if r.time_saved)
    avg_time_saved = total_time_saved / inefficient_routes if inefficient_routes else 0

    total_cost_saved = 0
    cost_table = []
    for r in routes:
        try:
            actual_duration = (r.actual_delivery_time - r.starting_time).total_seconds() / 3600.0
        except Exception:
            continue
        if r.optimized_delivery_time is None:
            continue

        optimized_time = r.optimized_delivery_time
        optimized_cost = r.max_delivery_cost * optimized_time
        actual_cost = r.max_delivery_cost * actual_duration
        cost_saved = actual_cost - optimized_cost
        total_cost_saved += cost_saved

        cost_table.append({
            "route_id": r.id,
            "base_address": r.base_address,
            "shipping_address": r.shipping_address,
            "actual_duration": round(actual_duration, 2),
            "optimized_time": round(optimized_time, 2),
            "max_delivery_cost": r.max_delivery_cost,
            "optimized_cost": round(optimized_cost, 2),
            "actual_cost": round(actual_cost, 2),
            "cost_saved": round(cost_saved, 2)
        })

    avg_cost_saved = total_cost_saved / inefficient_routes if inefficient_routes else 0

    ineff_table = []
    for r in routes:
        try:
            delay = (r.actual_delivery_time - r.expected_delivery_time).total_seconds() / 3600.0
        except Exception:
            delay = "N/A"
        ineff_table.append({
            "file_id": r.file_id,
            "base_address": r.base_address,
            "shipping_address": r.shipping_address,
            "starting_time": r.starting_time.strftime("%Y-%m-%d %H:%M"),
            "expected_delivery_time": r.expected_delivery_time.strftime("%Y-%m-%d %H:%M"),
            "actual_delivery_time": r.actual_delivery_time.strftime("%Y-%m-%d %H:%M"),
            "expected_delivery_cost": r.expected_delivery_cost,
            "actual_delivery_cost": r.actual_delivery_cost,
            "max_delivery_cost": r.max_delivery_cost,
            "delay_hours": round(delay, 2) if isinstance(delay, (int, float)) else delay,
            "optimized_delivery_time": r.optimized_delivery_time if r.optimized_delivery_time is not None else "N/A",
            "time_saved": r.time_saved if r.time_saved is not None else "N/A"
        })

    chart_url = generate_cost_waste_chart(cost_table, file_id)

    return {
        "file_name": file.filename,
        "upload_date": file.upload_date,
        "user_name": f"{user.first_name} {user.last_name}",
        "user_email": user.email,
        "inefficient_routes": inefficient_routes,
        "total_delayed_hours": round(total_delayed_hours, 2),
        "avg_delayed_hours": round(avg_delayed_hours, 2),
        "total_time_saved": round(total_time_saved, 2),
        "avg_time_saved": round(avg_time_saved, 2),
        "total_cost_saved": round(total_cost_saved, 2),
        "avg_cost_saved": round(avg_cost_saved, 2),
        "ineff_table": ineff_table,
        "cost_table": cost_table,
        "chart_url": chart_url
    }


# --- Flask routes for Report Preview and Downloads ---

@app.route('/report/<int:file_id>')
def report_preview(file_id):
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)
    chart_url = f"/static/plots/cost_waste_{file_id}.png"
    return render_template("report_preview.html", summary=summary, file_id=file_id, chart_url=chart_url)


@app.route('/download_report_pdf/<int:file_id>')
def download_report_pdf(file_id):
    if 'user_id' not in session:
        flash("Login required.", "danger")
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "RetroTrack", ln=True, align='C')
    pdf.cell(0, 10, "Logistics Inefficiency Report", ln=True, align='C')
    pdf.ln(10)

    pdf.set_font("Arial", '', 12)
    for key in ['file_name', 'upload_date', 'user_name', 'user_email',
                'inefficient_routes', 'total_delayed_hours', 'avg_delayed_hours',
                'total_time_saved', 'avg_time_saved', 'total_cost_saved', 'avg_cost_saved']:
        value = summary[key]
        pdf.cell(0, 8, f"{key.replace('_', ' ').title()}: {value}", ln=True)

    # Table for Inefficient Routes
    pdf.ln(8)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Inefficient Routes", ln=True)
    pdf.set_font("Arial", '', 10)
    for row in summary["ineff_table"]:
        pdf.multi_cell(0, 6, json.dumps(row, ensure_ascii=False))

    # Table for Cost Analysis
    pdf.ln(8)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Cost Analysis", ln=True)
    pdf.set_font("Arial", '', 10)
    for row in summary["cost_table"]:
        pdf.multi_cell(0, 6, json.dumps(row, ensure_ascii=False))

    # Chart
    chart_path = f"static/plots/cost_waste_{file_id}.png"
    if os.path.exists(chart_path):
        pdf.ln(10)
        pdf.image(chart_path, x=10, w=180)

    output = io.BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin-1')
    output.write(pdf_output)
    output.seek(0)

    return send_file(output,
                     mimetype='application/pdf',
                     as_attachment=True,
                     download_name=f'report_{file_id}.pdf')


@app.route('/download_report_word/<int:file_id>')
def download_report_word(file_id):
    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    doc = Document()
    doc.add_heading("RetroTrack", 0)
    doc.add_heading("Logistics Inefficiency Report", level=1)
    doc.add_paragraph(f"Generated by {summary['user_name']} ({summary['user_email']})")
    doc.add_paragraph("\n")

    for key in ['file_name', 'upload_date', 'inefficient_routes', 'total_delayed_hours',
                'avg_delayed_hours', 'total_time_saved', 'avg_time_saved',
                'total_cost_saved', 'avg_cost_saved']:
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {summary[key]}")

    doc.add_heading("Inefficient Routes", level=2)
    for row in summary["ineff_table"]:
        doc.add_paragraph(json.dumps(row, ensure_ascii=False))

    doc.add_heading("Cost Analysis", level=2)
    for row in summary["cost_table"]:
        doc.add_paragraph(json.dumps(row, ensure_ascii=False))

    chart_path = f"static/plots/cost_waste_{file_id}.png"
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     download_name=f"report_{file_id}.docx", as_attachment=True)


@app.route('/download_report_excel/<int:file_id>')
def download_report_excel(file_id):
    user = User.query.get(session['user_id'])
    summary = generate_summary(file_id, user)

    buf = io.BytesIO()
    workbook = xlsxwriter.Workbook(buf)
    summary_ws = workbook.add_worksheet("Summary")
    ineff_ws = workbook.add_worksheet("Inefficient Routes")
    cost_ws = workbook.add_worksheet("Cost Analysis")

    # Write summary
    row = 0
    for key in ['file_name', 'upload_date', 'user_name', 'user_email', 'inefficient_routes',
                'total_delayed_hours', 'avg_delayed_hours', 'total_time_saved',
                'avg_time_saved', 'total_cost_saved', 'avg_cost_saved']:
        summary_ws.write(row, 0, key.replace('_', ' ').title())
        summary_ws.write(row, 1, summary[key])
        row += 1

    # Write Inefficient Routes Table
    headers = list(summary['ineff_table'][0].keys()) if summary['ineff_table'] else []
    for col, h in enumerate(headers):
        ineff_ws.write(0, col, h)
    for i, item in enumerate(summary['ineff_table'], 1):
        for j, h in enumerate(headers):
            ineff_ws.write(i, j, item[h])

    # Write Cost Table
    headers = list(summary['cost_table'][0].keys()) if summary['cost_table'] else []
    for col, h in enumerate(headers):
        cost_ws.write(0, col, h)
    for i, item in enumerate(summary['cost_table'], 1):
        for j, h in enumerate(headers):
            cost_ws.write(i, j, item[h])

    # Insert chart
    chart_path = f"static/plots/cost_waste_{file_id}.png"
    if os.path.exists(chart_path):
        cost_ws.insert_image('N2', chart_path)

    workbook.close()
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     download_name=f"report_{file_id}.xlsx", as_attachment=True)


#########################################################################
# Main Entry
#########################################################################
if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
